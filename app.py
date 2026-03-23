import gradio as gr
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import textwrap, json, csv, os, re

VALID_KEYS = set(os.environ.get("raxzen_all_file_tool_key", "raxzen_file_create_key").split(","))

def verify_key(api_key: str) -> bool:
    return api_key.strip() in VALID_KEYS

def generate_pdf(text, title, font_size, page_size, bg_color, text_color, add_border):
    filepath = "/tmp/output.pdf"
    size = A4 if page_size == "A4" else LETTER
    doc = SimpleDocTemplate(filepath, pagesize=size,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()

    def hex_to_rgb(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16)/255 for i in (0, 2, 4))

    rc = colors.Color(*hex_to_rgb(text_color))
    title_style = ParagraphStyle("t", parent=styles["Title"],
                                 fontSize=int(font_size)+4, textColor=rc)
    body_style  = ParagraphStyle("b", parent=styles["Normal"],
                                 fontSize=int(font_size), textColor=rc,
                                 leading=int(font_size)*1.5)
    story = []
    if title.strip():
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
    for line in text.split("\n"):
        story.append(Paragraph(line if line.strip() else "&nbsp;", body_style))
        story.append(Spacer(1, 0.05*inch))

    def draw_bg(c, d):
        c.saveState()
        w, h = size
        br, bg, bb = hex_to_rgb(bg_color)
        c.setFillColorRGB(br, bg, bb)
        c.rect(0, 0, w, h, fill=1, stroke=0)
        if add_border:
            c.setStrokeColorRGB(*hex_to_rgb(text_color))
            c.setLineWidth(2)
            c.rect(20, 20, w-40, h-40, fill=0, stroke=1)
        c.restoreState()

    doc.build(story, onFirstPage=draw_bg, onLaterPages=draw_bg)
    return filepath

def generate_html(text, title, theme, font_family, custom_css):
    themes = {
        "Light": ("background:#f9f9f9;color:#222;", "#fff"),
        "Dark":  ("background:#1a1a2e;color:#eee;", "#16213e"),
        "Ocean": ("background:#e0f7fa;color:#006064;", "#fff"),
        "Sunset":("background:#fff3e0;color:#e65100;", "#fff"),
    }
    body_style, card_bg = themes.get(theme, themes["Light"])
    fonts = {"Default":"Arial,sans-serif","Serif":"Georgia,serif",
             "Mono":"'Courier New',monospace","Modern":"'Segoe UI',sans-serif"}
    font = fonts.get(font_family, "Arial,sans-serif")
    paras = "".join(f"<p>{l}</p>" if l.strip() else "<br>" for l in text.split("\n"))
    html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>{title or 'Document'}</title>
<style>
body{{font-family:{font};{body_style}margin:0;padding:20px}}
.card{{background:{card_bg};border-radius:10px;padding:30px;
       max-width:800px;margin:30px auto;box-shadow:0 4px 15px rgba(0,0,0,.1)}}
h1{{margin-top:0}}p{{line-height:1.8}}{custom_css}
</style></head><body><div class="card">
{"<h1>"+title+"</h1>" if title.strip() else ""}{paras}
</div></body></html>"""
    f = "/tmp/output.html"
    open(f,"w",encoding="utf-8").write(html)
    return f

def generate_txt(text, title, add_line_numbers, line_separator):
    sep = {"None":"","Dashes":"-"*60,"Equals":"="*60,"Stars":"*"*60}[line_separator]
    lines = text.split("\n")
    result = []
    if title.strip():
        result += [title.upper(), sep or "", ""]
    for i, l in enumerate(lines, 1):
        result.append(f"{i:>4}. {l}" if add_line_numbers else l)
    f = "/tmp/output.txt"
    open(f,"w",encoding="utf-8").write("\n".join(result))
    return f

def generate_docx(text, title, font_name, font_size, bold_title, alignment):
    align_map = {"Left":WD_ALIGN_PARAGRAPH.LEFT,"Center":WD_ALIGN_PARAGRAPH.CENTER,
                 "Right":WD_ALIGN_PARAGRAPH.RIGHT,"Justify":WD_ALIGN_PARAGRAPH.JUSTIFY}
    al = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    doc = Document()
    if title.strip():
        h = doc.add_heading(title, 1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs:
            r.bold = bold_title
            r.font.name = font_name
            r.font.size = Pt(int(font_size)+6)
    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        p.alignment = al
        for r in p.runs:
            r.font.name = font_name
            r.font.size = Pt(int(font_size))
    f = "/tmp/output.docx"
    doc.save(f)
    return f

def generate_csv_file(text, delimiter):
    delim = {"Comma":",","Tab":"\t","Semicolon":";","Pipe":"|"}[delimiter]
    f = "/tmp/output.csv"
    with open(f,"w",newline="",encoding="utf-8") as fh:
        csv.writer(fh, delimiter=delim).writerows([[l] for l in text.split("\n")])
    return f

def generate_json_file(text, title):
    data = {"title": title, "content": text.split("\n")}
    f = "/tmp/output.json"
    json.dump(data, open(f,"w",encoding="utf-8"), indent=2, ensure_ascii=False)
    return f

def generate_md(text, title):
    lines = ([f"# {title}\n"] if title.strip() else []) + text.split("\n")
    f = "/tmp/output.md"
    open(f,"w",encoding="utf-8").write("\n".join(lines))
    return f

def api_generate(api_key, text, title, file_type,
                 font_size=12, page_size="A4",
                 bg_color="#ffffff", text_color="#000000", add_border=False,
                 html_theme="Light", font_family="Default", custom_css="",
                 add_line_numbers=False, line_separator="None",
                 docx_font="Arial", bold_title=True, docx_alignment="Left",
                 csv_delimiter="Comma"):
    if not verify_key(api_key):
        return None, "❌ Invalid API Key"
    if not text.strip():
        return None, "❌ Text cannot be empty"
    try:
        ft = file_type
        if ft == "PDF":
            f = generate_pdf(text, title, font_size, page_size, bg_color, text_color, add_border)
        elif ft == "HTML":
            f = generate_html(text, title, html_theme, font_family, custom_css)
        elif ft == "TXT":
            f = generate_txt(text, title, add_line_numbers, line_separator)
        elif ft == "DOCX":
            f = generate_docx(text, title, docx_font, font_size, bold_title, docx_alignment)
        elif ft == "CSV":
            f = generate_csv_file(text, csv_delimiter)
        elif ft == "JSON":
            f = generate_json_file(text, title)
        elif ft == "Markdown":
            f = generate_md(text, title)
        else:
            return None, f"❌ Unknown type: {ft}"
        return f, f"✅ {ft} generated!"
    except Exception as e:
        return None, f"❌ Error: {str(e)}"

def api_read(api_key, file):
    if not verify_key(api_key):
        return "❌ Invalid API Key", ""
    if file is None:
        return "❌ No file uploaded", ""

    if isinstance(file, str):
        filepath = file
    elif hasattr(file, "name"):
        filepath = file.name
    else:
        return "❌ Cannot read file object", ""

    name = filepath.lower()

    def detect_type(path):
        try:
            with open(path, "rb") as f:
                header = f.read(8)
            if header[:4] == b"%PDF": return ".pdf"
            if header[:2] == b"PK":   return ".docx"
            with open(path, encoding="utf-8", errors="ignore") as f:
                sample = f.read(300).lower()
            if "<!doctype" in sample or "<html" in sample: return ".html"
            try:
                json.loads(open(path, encoding="utf-8", errors="ignore").read())
                return ".json"
            except: pass
            if "," in sample or ";" in sample: return ".csv"
            return ".txt"
        except:
            return ".txt"

    ext = ""
    for e in [".pdf",".docx",".html",".htm",".txt",".md",".csv",".json"]:
        if name.endswith(e):
            ext = e
            break
    if not ext:
        ext = detect_type(filepath)

    try:
        if ext == ".pdf":
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text, f"📄 PDF | {len(text)} chars"
        elif ext == ".docx":
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            return text, f"📝 DOCX | {len(text)} chars"
        elif ext in (".html", ".htm"):
            soup = BeautifulSoup(open(filepath, encoding="utf-8", errors="ignore"), "html.parser")
            text = soup.get_text("\n")
            return text, f"🌐 HTML | {len(text)} chars"
        elif ext in (".txt", ".md"):
            text = open(filepath, encoding="utf-8", errors="ignore").read()
            return text, f"📃 TXT | {len(text)} chars"
        elif ext == ".csv":
            text = open(filepath, encoding="utf-8", errors="ignore").read()
            return text, f"📊 CSV | {len(text.splitlines())} rows"
        elif ext == ".json":
            data = json.load(open(filepath, encoding="utf-8"))
            text = json.dumps(data, indent=2, ensure_ascii=False)
            return text, f"🔧 JSON | {len(text)} chars"
        return "❌ Unsupported file type", ""
    except Exception as e:
        return f"❌ Error: {str(e)}", ""

def api_edit(api_key, text, action,
             find_text="", replace_text="",
             prefix="", suffix="", wrap_width=80):
    if not verify_key(api_key):
        return "❌ Invalid API Key"
    if not text:
        return ""
    if action == "Add Prefix/Suffix":
        return "\n".join(f"{prefix}{l}{suffix}" for l in text.split("\n"))
    elif action == "Find & Replace":
        return text.replace(find_text, replace_text)
    elif action == "UPPERCASE":    return text.upper()
    elif action == "lowercase":    return text.lower()
    elif action == "Title Case":   return text.title()
    elif action == "Remove Extra Spaces":
        return re.sub(r" +", " ", text).strip()
    elif action == "Remove Empty Lines":
        return "\n".join(l for l in text.split("\n") if l.strip())
    elif action == "Word Wrap":
        return "\n".join(textwrap.fill(l,int(wrap_width)) for l in text.split("\n"))
    elif action == "Reverse Lines":  return "\n".join(reversed(text.split("\n")))
    elif action == "Sort Lines A-Z": return "\n".join(sorted(text.split("\n")))
    elif action == "Numbered Lines":
        return "\n".join(f"{i+1}. {l}" for i,l in enumerate(text.split("\n")))
    elif action == "Count Words":
        wc=len(text.split()); cc=len(text); lc=text.count("\n")+1
        return f"Words:{wc} | Chars:{cc} | Lines:{lc}"
    return text

def analyze_image_multi(api_key, image):
    if not verify_key(api_key):
        return "❌ Invalid API Key"
    if image is None:
        return "❌ No image uploaded"

    from PIL import Image
    import numpy as np
    from concurrent.futures import ThreadPoolExecutor, as_completed

    if isinstance(image, np.ndarray):
        img = Image.fromarray(image).convert("RGB")
    elif isinstance(image, str):
        img = Image.open(image).convert("RGB")
    else:
        img = image.convert("RGB")

    def run_blip2(img):
        from transformers import BlipProcessor, BlipForConditionalGeneration
        import torch
        processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-large")
        model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-large")
        model.eval()
        inputs = processor(img, return_tensors="pt")
        with torch.no_grad():
            out = model.generate(**inputs, max_new_tokens=120)
        caption = processor.decode(out[0], skip_special_tokens=True)
        return f"BLIP-2\nCaption: {caption}\nSize: {img.width}x{img.height}"

    def run_llava(img):
        from transformers import LlavaForConditionalGeneration, AutoProcessor
        import torch
        model_id = "llava-hf/llava-interleave-qwen-0.5b-hf"
        processor = AutoProcessor.from_pretrained(model_id)
        model = LlavaForConditionalGeneration.from_pretrained(
            model_id, torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
        )
        model.eval()
        prompt = "USER: <image>\nWhat is in this image? Describe in detail.\nASSISTANT:"
        inputs = processor(text=prompt, images=img, return_tensors="pt")
        with torch.no_grad():
            out = model.generate(**inputs, max_new_tokens=200)
        answer = processor.decode(out[0], skip_special_tokens=True)
        answer = answer.split("ASSISTANT:")[-1].strip()
        return f"LLaVA (Real)\nAnswer: {answer}"

    results = {}
    errors = []

    with ThreadPoolExecutor(max_workers=2) as ex:
        futures = {
            ex.submit(run_blip2, img): "BLIP-2",
            ex.submit(run_llava, img): "LLaVA",
        }
        for future in as_completed(futures):
            name = futures[future]
            try:
                results[name] = future.result()
            except Exception as e:
                errors.append(f"{name}: {str(e)[:80]}")

    sep = "\n\n" + "-"*40 + "\n\n"
    final = sep.join([results.get("BLIP-2",""), results.get("LLaVA","")]).strip(sep)
    if not final:
        final = "❌ All libraries failed"
    if errors:
        final += "\n\nErrors:\n" + "\n".join(f"  {e}" for e in errors)
    return final


def analyze_audio_multi(api_key, audio_file):
    if not verify_key(api_key):
        return "❌ Invalid API Key"
    if audio_file is None:
        return "❌ No audio uploaded"

    filepath = audio_file if isinstance(audio_file, str) else audio_file.name
    from concurrent.futures import ThreadPoolExecutor, as_completed

    def run_whisper(path):
        import whisper
        wmodel = whisper.load_model("base")
        res = wmodel.transcribe(path, task="transcribe")
        text = res["text"].strip()
        lang = res.get("language", "unknown")
        segs = res.get("segments", [])
        dur = segs[-1]["end"] if segs else 0
        return (
            "🥇 Whisper\n"
            f"Language: {lang.upper()}\n"
            f"Duration: {dur:.1f}s\n"
            f"Transcript: {text[:600] if text else '(no speech detected)'}"
        )

    def run_wav2vec2(path):
        import torch, soundfile as sf
        from transformers import Wav2Vec2Processor, Wav2Vec2ForCTC
        processor = Wav2Vec2Processor.from_pretrained("facebook/wav2vec2-base-960h")
        model = Wav2Vec2ForCTC.from_pretrained("facebook/wav2vec2-base-960h")
        model.eval()
        speech, sr = sf.read(path)
        # mono করো যদি stereo হয়
        if len(speech.shape) > 1:
            speech = speech.mean(axis=1)
        # 16kHz resample দরকার হলে
        if sr != 16000:
            import librosa
            speech = librosa.resample(speech, orig_sr=sr, target_sr=16000)
        inputs = processor(speech, sampling_rate=16000, return_tensors="pt", padding=True)
        with torch.no_grad():
            logits = model(**inputs).logits
        predicted_ids = torch.argmax(logits, dim=-1)
        transcript = processor.batch_decode(predicted_ids)[0]
        return (
            "🥈 Wav2Vec2 (Deep Speech AI)\n"
            f"Transcript: {transcript[:600] if transcript.strip() else '(no speech detected)'}"
        )

    results = {}
    errors = []

    with ThreadPoolExecutor(max_workers=2) as ex:
        futures = {
            ex.submit(run_whisper,  filepath): "Whisper",
            ex.submit(run_wav2vec2, filepath): "Wav2Vec2",
        }
        for future in as_completed(futures):
            name = futures[future]
            try:
                results[name] = future.result()
            except Exception as e:
                errors.append(f"{name}: {str(e)[:100]}")

    sep = "\n\n" + "-"*40 + "\n\n"
    parts = [r for r in [results.get("Whisper",""), results.get("Wav2Vec2","")] if r]
    final = sep.join(parts) if parts else "❌ All libraries failed"
    if errors:
        final += "\n\nErrors:\n" + "\n".join(f"  ⚠️ {e}" for e in errors)
    return final


CSS = "footer{display:none!important}"

with gr.Blocks(title="⚡ AllFile Tool Pro") as app:

    gr.Markdown("# ⚡ AllFile Tool Pro\n**PDF · HTML · TXT · DOCX · CSV · JSON · Markdown · 🖼️ Image · 🔊 Audio**")

    with gr.Tab("📄 Generate"):
        with gr.Row():
            with gr.Column(scale=2):
                g_key   = gr.Textbox(label="🔑 API Key", type="password",
                                     placeholder="Enter your API key")
                g_title = gr.Textbox(label="Title", placeholder="My Document")
                g_text  = gr.Textbox(label="Content", lines=10,
                                     placeholder="Type content here...")
                g_type  = gr.Dropdown(["PDF","HTML","TXT","DOCX","CSV","JSON","Markdown"],
                                      value="PDF", label="File Type")
            with gr.Column(scale=1):
                g_fs    = gr.Slider(8, 32, value=12, step=1, label="Font Size")
                g_ps    = gr.Radio(["A4","Letter"], value="A4", label="Page Size")
                g_bg    = gr.ColorPicker(value="#ffffff", label="BG Color")
                g_tc    = gr.ColorPicker(value="#000000", label="Text Color")
                g_brd   = gr.Checkbox(label="Add Border")
                g_theme = gr.Dropdown(["Light","Dark","Ocean","Sunset"], value="Light", label="HTML Theme")
                g_font  = gr.Dropdown(["Default","Serif","Mono","Modern"], value="Default", label="Font Family")
                g_css   = gr.Textbox(label="Custom CSS", lines=2)
                g_ln    = gr.Checkbox(label="Line Numbers (TXT)")
                g_sep   = gr.Dropdown(["None","Dashes","Equals","Stars"], value="None", label="Separator")
                g_df    = gr.Dropdown(["Arial","Times New Roman","Calibri"], value="Arial", label="DOCX Font")
                g_db    = gr.Checkbox(label="Bold Title", value=True)
                g_da    = gr.Dropdown(["Left","Center","Right","Justify"], value="Left", label="Alignment")
                g_cd    = gr.Dropdown(["Comma","Tab","Semicolon","Pipe"], value="Comma", label="CSV Delimiter")

        g_btn    = gr.Button("🚀 Generate", variant="primary")
        g_status = gr.Textbox(label="Status", interactive=False)
        g_out    = gr.File(label="⬇️ Download")

        g_btn.click(api_generate,
                    inputs=[g_key,g_text,g_title,g_type,
                            g_fs,g_ps,g_bg,g_tc,g_brd,
                            g_theme,g_font,g_css,
                            g_ln,g_sep,g_df,g_db,g_da,g_cd],
                    outputs=[g_out,g_status])

    with gr.Tab("📖 Read"):
        r_key  = gr.Textbox(label="🔑 API Key", type="password")
        r_file = gr.File(label="Upload File")
        r_btn  = gr.Button("📖 Read", variant="primary")
        r_info = gr.Textbox(label="Info", interactive=False)
        r_out  = gr.Textbox(label="Content", lines=15, interactive=True)
        r_btn.click(api_read, inputs=[r_key,r_file], outputs=[r_out,r_info])

    with gr.Tab("✏️ Edit"):
        e_key  = gr.Textbox(label="🔑 API Key", type="password")
        with gr.Row():
            with gr.Column():
                e_in  = gr.Textbox(label="Input Text", lines=10)
                e_act = gr.Dropdown([
                    "Find & Replace","Add Prefix/Suffix","UPPERCASE","lowercase",
                    "Title Case","Remove Extra Spaces","Remove Empty Lines",
                    "Word Wrap","Reverse Lines","Sort Lines A-Z",
                    "Numbered Lines","Count Words"
                ], value="Find & Replace", label="Action")
            with gr.Column():
                e_find  = gr.Textbox(label="Find")
                e_repl  = gr.Textbox(label="Replace")
                e_pre   = gr.Textbox(label="Prefix")
                e_suf   = gr.Textbox(label="Suffix")
                e_wrap  = gr.Number(label="Wrap Width", value=80)
        e_btn = gr.Button("⚡ Apply", variant="primary")
        e_out = gr.Textbox(label="Result", lines=10)
        e_btn.click(api_edit,
                    inputs=[e_key,e_in,e_act,e_find,e_repl,e_pre,e_suf,e_wrap],
                    outputs=e_out)

    with gr.Tab("🖼️ Image Analysis"):
        gr.Markdown("## 🖼️ Image Analysis\n🥇 BLIP-2 + 🥉 YOLOv8 + 🏅 CLIP + 🎖️ OCR — সব একসাথে চলবে")
        with gr.Row():
            with gr.Column(scale=1):
                img_key    = gr.Textbox(label="🔑 API Key", type="password",
                                        placeholder="Enter your API key")
                img_upload = gr.Image(label="📷 Upload Image",type="numpy")
                img_btn    = gr.Button("🔍 Analyze (All Libraries)", variant="primary")
            with gr.Column(scale=1):
                img_out = gr.Textbox(label="📊 Combined Result",
                                     lines=25, interactive=False)
        img_btn.click(
            analyze_image_multi,
            inputs=[img_key, img_upload],
            outputs=img_out,
            api_name="analyze_image"
        )

    with gr.Tab("🔊 Audio Analysis"):
        gr.Markdown("## 🔊 Audio Analysis\n🥇 Whisper + 🏅 Librosa + 🎖️ PyDub — সব একসাথে চলবে")
        with gr.Row():
            with gr.Column(scale=1):
                aud_key    = gr.Textbox(label="🔑 API Key", type="password",
                                        placeholder="Enter your API key")
                aud_upload = gr.Audio(label="🎵 Upload Audio / Song / Sound",
                                      type="filepath")
                aud_btn    = gr.Button("🎧 Analyze (All Libraries)", variant="primary")
            with gr.Column(scale=1):
                aud_out = gr.Textbox(label="📊 Combined Result",
                                     lines=25, interactive=False)
        aud_btn.click(
            analyze_audio_multi,
            inputs=[aud_key, aud_upload],
            outputs=aud_out,
            api_name="analyze_audio"
        )

    with gr.Tab("📡 API Docs"):
        gr.Markdown("""

**তোমার Space URL:**
```
https://rafidmondal-allfile-tool.hf.space
```

---

```html
<script type="module">
import { Client } from "https://cdn.jsdelivr.net/npm/@gradio/client/dist/index.min.js";

async function generateFile() {
  const client = await Client.connect("Rafidmondal/allfile-tool");

  const result = await client.predict("/api_generate", {
    api_key: "raxzen_file_create_key",
    text: "Hello from my website!",
    title: "My Document",
    file_type: "PDF",
    font_size: 12,
    page_size: "A4",
    bg_color: "#ffffff",
    text_color: "#000000",
    add_border: false,
    html_theme: "Light",
    font_family: "Default",
    custom_css: "",
    add_line_numbers: false,
    line_separator: "None",
    docx_font: "Arial",
    bold_title: true,
    docx_alignment: "Left",
    csv_delimiter: "Comma"
  });

  // result.data = [file_url, status_message]
  const fileUrl = result.data[0].url;
  window.open(fileUrl, "_blank"); // auto download!
}
</script>
```

---

```python
from gradio_client import Client

client = Client("Rafidmondal/allfile-tool")

result = client.predict(
    api_key="raxzen_file_create_key",
    text="Hello World",
    title="My Doc",
    file_type="PDF",
    font_size=12,
    page_size="A4",
    bg_color="#ffffff",
    text_color="#000000",
    add_border=False,
    html_theme="Light",
    font_family="Default",
    custom_css="",
    add_line_numbers=False,
    line_separator="None",
    docx_font="Arial",
    bold_title=True,
    docx_alignment="Left",
    csv_delimiter="Comma",
    api_name="/api_generate"
)

file_path, status = result
print(status)   # ✅ PDF generated!
print(file_path) # /tmp/output.pdf
```

---

---

```html
<script type="module">
import { Client } from "https://cdn.jsdelivr.net/npm/@gradio/client/dist/index.min.js";

async function analyzeImage(imageFile) {
  const client = await Client.connect("Rafidmondal/allfile-tool");

  const result = await client.predict("/analyze_image", {
    api_key: "raxzen_file_create_key",
    image: imageFile,   // File object (from <input type="file">)
    library: "🥇 BLIP-2 (~85-90%) — Caption & Understanding"
    // অন্য options:
    // "🥈 LLaVA (~85-90%) — Scene Explanation"
    // "🥉 YOLOv8 (~80-92%) — Object Detection"
    // "🏅 CLIP (~80-88%) — Image Classification"
    // "🎖️ Tesseract OCR (~60-85%) — Text Extraction"
  });

  console.log(result.data[0]); // analysis result text
}
</script>
```

```python
from gradio_client import Client, handle_file

client = Client("Rafidmondal/allfile-tool")

result = client.predict(
    api_key="raxzen_file_create_key",
    image=handle_file("photo.jpg"),
    library="🥇 BLIP-2 (~85-90%) — Caption & Understanding",
    api_name="/analyze_image"
)
print(result)  # 📝 Caption: a dog sitting on grass ...
```

---

```html
<script type="module">
import { Client } from "https://cdn.jsdelivr.net/npm/@gradio/client/dist/index.min.js";

async function analyzeAudio(audioFile) {
  const client = await Client.connect("Rafidmondal/allfile-tool");

  const result = await client.predict("/analyze_audio", {
    api_key: "raxzen_file_create_key",
    audio_file: audioFile,  // File object
    library: "🥇 Whisper (~95-100%) — Speech to Text"
    // অন্য options:
    // "🥉 Wav2Vec2 (~85-95%) — Deep Speech AI"
    // "🏅 Librosa (~75-85%) — Music Analysis"
    // "🎖️ PyDub (~70-80%) — Audio Info & Edit"
  });

  console.log(result.data[0]); // transcription / analysis
}
</script>
```

```python
from gradio_client import Client, handle_file

client = Client("Rafidmondal/allfile-tool")

result = client.predict(
    api_key="raxzen_file_create_key",
    audio_file=handle_file("song.mp3"),
    library="🏅 Librosa (~75-85%) — Music Analysis",
    api_name="/analyze_audio"
)
print(result)  # 🥁 Tempo: 128.0 BPM | 🎹 Key: A ...
```

---

`Settings → Variables and secrets → New secret`
- **Name:** `API_KEYS`
- **Value:** `mykey1,mykey2,mykey3`

---

```
gradio
reportlab
pdfplumber
python-docx
beautifulsoup4
transformers
torch
torchvision
Pillow
ultralytics
pytesseract
openai-whisper
librosa
pydub
soundfile
ffmpeg-python
```
        """)

app.launch(server_name="0.0.0.0", css=CSS, theme=gr.themes.Soft())
